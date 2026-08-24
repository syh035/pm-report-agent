# -*- coding: utf-8 -*-
"""
周报导出工具：把生成的 HTML 周报导出为 Word 和纯文本。

- html_to_docx(report_html, out_path)  -> python-docx 生成 Word（标题/表格/加粗/颜色）
- html_to_text(report_html)            -> 用 html.parser 提取纯文本
"""
from __future__ import annotations
import io
import re
from html.parser import HTMLParser
from typing import List, Tuple


# ---------- HTML -> 纯文本 ----------
class _TextExtractor(HTMLParser):
    """从 HTML 提取纯文本，保留换行与表格结构。"""
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts: List[str] = []
        self._tag_stack: List[str] = []
        self._in_table = False
        self._in_td = False
        self._last_space = False

    def handle_starttag(self, tag, attrs):
        self._tag_stack.append(tag)
        if tag in ("h1", "h2", "h3", "h4", "p", "div", "li", "tr"):
            self._newline()
        elif tag == "table":
            self._in_table = True
            self._newline()
        elif tag == "tr":
            self._newline()
        elif tag in ("th", "td"):
            self._in_td = True

    def handle_endtag(self, tag):
        if tag in self._tag_stack:
            while self._tag_stack:
                t = self._tag_stack.pop()
                if t == tag:
                    break
        if tag in ("h1", "h2", "h3", "p", "div", "li"):
            self._newline()
        elif tag == "tr":
            self._newline()
        elif tag in ("th", "td"):
            self.parts.append("  ")
            self._in_td = False
        elif tag == "table":
            self._in_table = False
            self._newline()

    def handle_data(self, data):
        text = re.sub(r"\s+", " ", data).strip()
        if text:
            if self.parts and self.parts[-1].endswith("  ") and not text.startswith(" "):
                pass
            self.parts.append(text)

    def _newline(self):
        if self.parts and self.parts[-1] != "\n":
            # 避免重复换行
            if self.parts[-1].endswith("  "):
                self.parts.append("\n")
            else:
                self.parts.append("\n")

    def get_text(self) -> str:
        out = "".join(self.parts)
        # 清理：去掉 td 后多余空格导致的换行问题
        out = re.sub(r"\n\s*\n+", "\n", out)
        out = re.sub(r"[ \t]+\n", "\n", out)
        return out.strip()


def html_to_text(html: str) -> str:
    parser = _TextExtractor()
    parser.feed(html or "")
    parser.close()
    return parser.get_text()


# ---------- HTML -> Word ----------
# 简单的 HTML → docx 转换器（覆盖标题、段落、表格、加粗、颜色）
def _parse_inline(paragraph, text_runs: List[Tuple[str, dict]]):
    """把简单行内 token 处理为 (text, style) 列表。这里简化为纯文本+粗略加粗。"""
    # 简化：text_runs 已在更上层处理，此函数保留占位
    return text_runs


def _strip_tags(text: str) -> str:
    """去掉行内 HTML 标签，返回纯文本。"""
    return re.sub(r"<[^>]+>", "", text or "")


def html_to_docx(html: str, out_path: str) -> str:
    """把 HTML 周报转成 Word 文档，保存到 out_path，返回路径。"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 页面横向留白适中
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    lines = _split_top_level(html)
    for block in lines:
        kind, content = block
        text = _strip_tags(content).strip()
        if not text:
            continue
        if kind == "h1":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor(0x2D, 0x31, 0x42)
            p.paragraph_format.space_after = Pt(6)
        elif kind == "h2":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x4A, 0x6C, 0xF7)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif kind == "p":
            _add_paragraph_with_inline(doc, content)
        elif kind == "li":
            _add_paragraph_with_inline(doc, content, bullet=True)
        elif kind == "table":
            _add_table(doc, content)
        # 其它块忽略

    doc.save(out_path)
    return out_path


def _split_top_level(html: str):
    """把 HTML 切成顶层块：(kind, inner_html)。支持 p/h1-h4/li/table。"""
    # 提取顶层块（方法：按标签切）
    blocks = []
    i = 0
    pos = 0
    pattern = re.compile(r"<(/?)(h[1-4]|p|li|table|ul|ol|div)\b[^>]*>", re.I)
    stack = []
    cur_kind = None
    cur_start = None

    def flush(kind, start, end):
        if start is not None and end is not None and end > start:
            inner = html[start:end]
            blocks.append((kind, inner))

    for m in pattern.finditer(html):
        closing, tag = m.group(1).lower(), m.group(2).lower()
        # 是否影响顶层分块：只关注 h1-h4,p,li,table,ul,ol
        if tag in ("ul", "ol"):
            continue  # 列表容器不单独成块
        if not closing:
            # 开始标签
            if tag in ("h1", "h2", "h3", "h4", "p", "li", "table"):
                flush(cur_kind, cur_start, m.start())
                cur_kind = tag
                cur_start = m.end()
        else:
            # 结束标签
            if tag == cur_kind:
                flush(cur_kind, cur_start, m.start())
                cur_kind = None
                cur_start = None
    flush(cur_kind, cur_start, None)
    return blocks


def _add_paragraph_with_inline(doc, html_text: str, bullet: bool = False):
    """把一个含 <b>/<span style=color...> 的段落加进 doc。"""
    from docx.shared import RGBColor, Pt
    from docx import Document
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    # 极简解析：按 <b> 和 <span style="color:#.."> 分段
    segs = _tokenize_inline(html_text)
    for text, bold, color in segs:
        if not text:
            continue
        run = p.add_run(text)
        if bold:
            run.bold = True
        if color:
            try:
                h = color.lstrip("#")
                run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
            except Exception:
                pass


def _tokenize_inline(html_text: str):
    """解析行内的 <b>、<span style="color:#xxx"> 等，返回 [(text, bold, color)]。"""
    tokens = []
    pattern = re.compile(
        r"<(/?)(b|strong|span)\b([^>]*)>",
        re.I,
    )
    pos = 0
    bold = False
    color = None
    for m in pattern.finditer(html_text):
        closing, tag, attrs = m.group(1).lower(), m.group(2).lower(), m.group(3)
        text = html_text[pos:m.start()]
        if text:
            tokens.append((re.sub(r"\s+", " ", text).strip(" "), bold, color))
        if tag in ("b", "strong"):
            bold = not closing
        elif tag == "span":
            if not closing:
                cm = re.search(r"color\s*:\s*#[0-9a-fA-F]{6}", attrs) or \
                     re.search(r"color\s*:\s*#[0-9a-fA-F]{3}", attrs)
                color = cm.group(0).split(":")[1].strip() if cm else color
            else:
                color = None
        pos = m.end()
    tail = html_text[pos:]
    if tail:
        tokens.append((re.sub(r"\s+", " ", tail).strip(" "), bold, color))
    return tokens


def _add_table(doc, inner_html: str):
    """把 <table> 内部转换为 docx 表格。"""
    from docx.shared import RGBColor, Pt, Cm
    rows = _extract_table_rows(inner_html)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Light Grid Accent 1"
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for c_i in range(n_cols):
            cell_text = _strip_tags(row[c_i]) if c_i < len(row) else ""
            cell = cells[c_i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.size = Pt(10.5)
            if r_i == 0:
                run.bold = True
                for cell_obj in cells:
                    for p_ in cell_obj.paragraphs:
                        for r_ in p_.runs:
                            r_.bold = True
                # 表头底色
    # 列宽自适应
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(4)


def _extract_table_rows(table_html: str) -> List[List[str]]:
    rows = []
    for tr_m in re.finditer(r"<tr[^>]*>(.*?)</tr>", table_html, re.S | re.I):
        tds = re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", tr_m.group(1), re.S | re.I)
        rows.append(tds)
    return rows


# 兼容：bytes 输出给 FastAPI
def html_to_docx_bytes(html: str) -> io.BytesIO:
    from docx import Document
    import os
    tmp = "/tmp/_pm_export.docx"
    html_to_docx(html, tmp)
    data = io.BytesIO(open(tmp, "rb").read())
    os.remove(tmp)
    return data
