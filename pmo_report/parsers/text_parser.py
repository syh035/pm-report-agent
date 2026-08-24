# -*- coding: utf-8 -*-
"""
文本解析器：Word (.docx) 和 PDF (.pdf)。

策略分两层：
  1) 规则层：用正则从文本里提取带数字/百分比/待办的段落，初步切成候选任务。
  2) AI 层（可选）：若配了 DeepSeek，用它把文本叙述结构化提炼成任务。
实现时先做规则层，AI 层通过 generator 注入，避免循环依赖。
"""
from __future__ import annotations
import os
import re
from typing import List

from ..models import (Project, Task, STATUS_DONE, STATUS_IN_PROGRESS,
                      STATUS_NOT_STARTED, STATUS_DELAYED, STATUS_RISK)
from ._date_util import parse_date


# ---------- 文本抽取 ----------

def _extract_text_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # 表格内容也纳入
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text_pdf(path: str, use_ocr: bool = True) -> str:
    import fitz  # PyMuPDF
    text = []
    doc = fitz.open(path)
    for page in doc:
        t = page.get_text()
        if t.strip():
            text.append(t)
    doc.close()
    joined = "\n".join(text)
    if joined.strip():
        return joined
    # 无文本层：走 OCR（复用 tesseract，路径按本机配置）
    if use_ocr:
        try:
            import pytesseract
            from pdf2image import convert_from_path
            pytesseract.pytesseract.tesseract_cmd = "/opt/homebrew/bin/tesseract"
            images = convert_from_path(path, dpi=200, poppler_path="/opt/homebrew/bin")
            ocr_parts = []
            for img in images:
                ocr_parts.append(pytesseract.image_to_string(img, lang="chi_sim+eng"))
            return "\n".join(ocr_parts)
        except Exception as e:
            joined = f"[OCR 不可用: {e}]"
    return joined


# ---------- 规则解析 ----------

# 段落特征：数字百分比 / 完成/进行/待 / 负责人 + 任务
_BULLET = re.compile(r"^\s*[•·▪▸►◦\-*—>\d.\u2460-\u2473\uff08)]")


_OWNER_PATTERNS = [
    # "xx 张三 负责" / "xx-张三" / "张三负责" / "负责人张三" 等
    re.compile(r"(?:负责人|归属|责任)\s*[:：]?\s*([\u4e00-\u9fa5]{2,4})"),
    re.compile(r"([\u4e00-\u9fa5]{2,4})\s*(?:负责|牵头|主管)"),
    re.compile(r"([\u4e00-\u9fa5]{2,4})\s+负责\s*[,，]?"),
]

# 切分"任务名"的断点：遇到 负责人/进度/状态词 前的纯名词作为任务名
_SPLIT_BY = [
    "负责", "牵头", "主管", "当前进度", "进度", "已完成", "完成", "进行中",
    "未开始", "有风险", "风险", "待开始", "即将", "预计", "将于", "计划",
]


def _extract_owner(line: str) -> str:
    for pat in _OWNER_PATTERNS:
        m = pat.search(line)
        if m:
            return m.group(0) if "负责人" in pat.pattern else m.group(1)
    return ""


def _extract_name(line: str) -> str:
    """从一行里剥离出任务名（去掉序号、负责人、状态、进度等干扰）。"""
    # 去序号
    name = _BULLET.sub("", line).strip()
    # 在第一个断点处截断
    for token in _SPLIT_BY:
        idx = name.find(token)
        if idx > 0:
            name = name[:idx].strip()
            break
    # 去掉可能残留的负责人名（"任务名 张三" 尾部的中文人名）
    m_owner_tail = re.search(r"\s([\u4e00-\u9fa5]{2,4})\s*$", name)
    if m_owner_tail:
        cand = name[:m_owner_tail.start()].strip()
        # 如果名字还有明显任务词，就砍掉这个疑似人名
        if re.search(r"(开发|模块|编写|部署|验收|测试|评审|设计|搭建|文档|接口|需求|系统|上线|联调|重构|报表|优化|升级|迁移|接入|改造)", cand):
            name = cand
    # 去掉可能残留的冒号/空格
    name = re.sub(r"[:：]\s*$", "", name).strip()
    return name[:50]


def _rule_parse(text: str, ignore: List[str] | None = None) -> tuple:
    """用规则把文本切成候选任务。主要识别带进度百分比的条目。
    返回 (tasks, ambiguous)：
      tasks     —— 命中的任务
      ambiguous —— 有任务特征但证据不足被过滤的行（供 AI 兜底，见 parse_docx/parse_pdf）
    ignore     —— 忽略词列表（规则库 ignore_keywords），命中则整行跳过"""
    ignore = ignore or []
    tasks: List[Task] = []
    ambiguous: List[str] = []
    seen = set()

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if any(kw in line for kw in ignore):   # 用户忽略词：整行跳过
            continue

        m_pct = re.search(r"(\d{1,3})\s*%", line)
        m_date = re.search(r"(\d{4}[-/]\d{1,2}[-/]\d{1,2})", line)
        m_status = re.search(r"(已完成|完成|已上线|进行中|未开始|待开始|有风险|风险|已滞后|已交付)", line)
        # 带有一个动作状态的词，更像任务条目（过滤纯标题/段落）
        has_action = bool(re.search(r"(负责|牵头|进度|完成|开发|编写|上线|部署|验收|测试|评审|推进|启动|完成度)", line))
        is_bullet = bool(_BULLET.match(line))

        if is_bullet:
            # 项目符号行：有百分比/状态词/动作词即视为任务条目
            need = m_pct or m_status or has_action
        else:
            # 非项目符号行：必须有"硬证据"（百分比/日期/负责人模式），
            # 防止"下周重点：完成联调并准备上线发布"这类纯叙述被误抓
            need = m_pct or m_date or bool(_extract_owner(line))

        if (is_bullet or m_pct or m_date or m_status) and need:
            owner = _extract_owner(line)
            name = _extract_name(line)
            if not name or name in seen:
                continue
            seen.add(name)

            t = Task(name=name, owner=owner, note=line)
            if m_pct:
                t.progress = min(int(m_pct.group(1)), 100)
                if t.progress >= 100:
                    t.status = STATUS_DONE
                elif t.progress > 0:
                    t.status = STATUS_IN_PROGRESS
                else:
                    t.status = STATUS_NOT_STARTED
            if m_status and not t.status:
                s = m_status.group(1)
                if s in ("已完成", "完成", "已上线", "已交付"):
                    t.status = STATUS_DONE
                elif s in ("进行中",):
                    t.status = STATUS_IN_PROGRESS
                elif s in ("未开始", "待开始"):
                    t.status = STATUS_NOT_STARTED
                elif s in ("已滞后",):
                    t.status = STATUS_DELAYED
                elif s in ("有风险", "风险"):
                    t.status = STATUS_RISK
            if m_date:
                d = parse_date(m_date.group(1))
                t.plan_end = d or t.plan_end
            tasks.append(t)
        elif has_action or m_status or m_pct:
            # 有任务特征但证据不足：记为模糊行，供 AI 兜底（不花规则层成本）
            if line not in ambiguous:
                ambiguous.append(line)
    return tasks, ambiguous


def _candidate_text(tasks: List[Task], max_chars: int = 3000) -> str:
    """把规则解析出的候选任务行拼成精简文本，供 AI 提炼使用。
    只送"像任务的行"，不送整篇原文 —— 省 Token 且降低噪音。"""
    lines = []
    for t in tasks:
        src = (t.note or t.name or "").strip()
        if src and src not in lines:
            lines.append(src)
    return "\n".join(lines)[:max_chars]


def _ai_candidate_text(tasks: List[Task], ambiguous: List[str], max_chars: int = 3000) -> str:
    """AI 兜底的候选文本 = 规则命中行 + 模糊行（仅这些行，控制 Token）。"""
    lines = []
    for t in tasks:
        src = (t.note or t.name or "").strip()
        if src and src not in lines:
            lines.append(src)
    for a in ambiguous:
        if a not in lines:
            lines.append(a)
    return "\n".join(lines)[:max_chars]


def parse_docx(path: str, project_name: str = "", use_ai: bool = True) -> Project:
    text = _extract_text_docx(path)
    from ..rules import custom_ignore_keywords
    tasks, ambiguous = _rule_parse(text, ignore=custom_ignore_keywords())
    proj = Project(name=project_name, tasks=tasks)

    # AI 提炼增强（可选）：规则命中行 + 模糊行送 AI（仅模糊行成本可控），失败保留规则结果
    if use_ai and (tasks or ambiguous):
        try:
            from ..ai import enrich_tasks_from_text
            cand = _ai_candidate_text(tasks, ambiguous)
            if cand:
                enriched = enrich_tasks_from_text(cand, tasks)
                if enriched:
                    proj.tasks = enriched
        except Exception:
            pass  # AI 失败则保留规则结果
    return proj


def parse_pdf(path: str, project_name: str = "", use_ai: bool = True) -> Project:
    text = _extract_text_pdf(path)
    from ..rules import custom_ignore_keywords
    tasks, ambiguous = _rule_parse(text, ignore=custom_ignore_keywords())
    proj = Project(name=project_name, tasks=tasks)
    if use_ai and (tasks or ambiguous):
        try:
            from ..ai import enrich_tasks_from_text
            cand = _ai_candidate_text(tasks, ambiguous)
            if cand:
                enriched = enrich_tasks_from_text(cand, tasks)
                if enriched:
                    proj.tasks = enriched
        except Exception:
            pass
    return proj
